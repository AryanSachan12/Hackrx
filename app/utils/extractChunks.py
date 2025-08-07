"""
Document Processor (Ingestion Pipeline)
Complete document processing system for extracting, structuring, and chunking documents.

Supports: PDF, DOCX, EML files
Features: URL downloads, structure preservation, semantic chunking
"""

# Document Processor (Ingestion Pipeline)
# Libraries and Imports
import os
import requests
import tempfile
from pathlib import Path
from typing import List, Dict, Any, Optional
from urllib.parse import urlparse, urljoin
import email
from email import policy
import mimetypes

# PDF Processing
import fitz  # PyMuPDF
import pdfplumber

# DOCX Processing  
from docx import Document

# Text Processing and Chunking
from langchain.text_splitter import RecursiveCharacterTextSplitter

# HTML Processing
from bs4 import BeautifulSoup


class DocumentDownloader:
    """Handles downloading files from URLs with proper error handling"""
    
    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        })
    
    def download_file(self, url: str, save_path: Optional[str] = None) -> str:
        """
        Download file from URL and return local file path
        
        Args:
            url: URL of the file to download
            save_path: Optional local path to save file
            
        Returns:
            str: Local file path of downloaded file
        """
        try:
            # Parse URL to get filename
            parsed_url = urlparse(url)
            filename = os.path.basename(parsed_url.path)
            
            # If no filename in URL, try to get from Content-Disposition header
            if not filename or '.' not in filename:
                response = self.session.head(url)
                content_disposition = response.headers.get('content-disposition', '')
                if 'filename=' in content_disposition:
                    filename = content_disposition.split('filename=')[1].strip('"\'')
                else:
                    # Generate filename based on content type
                    content_type = response.headers.get('content-type', '')
                    ext = mimetypes.guess_extension(content_type.split(';')[0]) or '.bin'
                    filename = f"document_{hash(url) % 10000}{ext}"
            
            # Determine save path
            if save_path is None:
                save_path = os.path.join(tempfile.gettempdir(), filename)
            
            # Download file
            print(f"📥 Downloading: {url}")
            response = self.session.get(url, stream=True)
            response.raise_for_status()
            
            # Save file
            os.makedirs(os.path.dirname(save_path), exist_ok=True)
            with open(save_path, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    f.write(chunk)
            
            print(f"✅ Downloaded to: {save_path}")
            return save_path
            
        except Exception as e:
            print(f"❌ Error downloading {url}: {str(e)}")
            raise
    
    def is_url(self, path: str) -> bool:
        """Check if string is a valid URL"""
        try:
            result = urlparse(path)
            return all([result.scheme, result.netloc])
        except:
            return False


class PDFProcessor:
    """Extract text from PDFs while preserving structure"""
    
    def __init__(self):
        self.pymupdf_available = True
        self.pdfplumber_available = True
    
    def extract_with_pymupdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text using PyMuPDF (fitz) with structure preservation"""
        try:
            doc = fitz.open(file_path)
            extracted_data = {
                'text': '',
                'metadata': {},
                'structure': {
                    'headings': [],
                    'tables': [],
                    'images': []
                }
            }
            
            # Extract metadata
            extracted_data['metadata'] = doc.metadata
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Extract text blocks with formatting info
                blocks = page.get_text("dict")
                
                for block in blocks["blocks"]:
                    if "lines" in block:  # Text block
                        for line in block["lines"]:
                            line_text = ""
                            for span in line["spans"]:
                                text = span["text"]
                                font_size = span["size"]
                                font_flags = span["flags"]
                                
                                # Detect headings based on font size and formatting
                                if font_size > 14 or font_flags & 2**4:  # Bold flag
                                    extracted_data['structure']['headings'].append({
                                        'text': text.strip(),
                                        'page': page_num + 1,
                                        'font_size': font_size
                                    })
                                    line_text += f"\n## {text}\n"
                                else:
                                    line_text += text
                            
                            extracted_data['text'] += line_text + "\n"
                
                # Extract tables (basic detection)
                tables = page.find_tables()
                for table in tables:
                    table_data = table.extract()
                    extracted_data['structure']['tables'].append({
                        'data': table_data,
                        'page': page_num + 1
                    })
                    
                    # Add table to text representation
                    if table_data:
                        extracted_data['text'] += "\n### Table\n"
                        for row in table_data:
                            if row:
                                extracted_data['text'] += " | ".join(str(cell) if cell else "" for cell in row) + "\n"
                        extracted_data['text'] += "\n"
            
            doc.close()
            return extracted_data
            
        except Exception as e:
            print(f"❌ PyMuPDF extraction failed: {str(e)}")
            return None
    
    def extract_with_pdfplumber(self, file_path: str) -> Dict[str, Any]:
        """Extract text using pdfplumber with table detection"""
        try:
            extracted_data = {
                'text': '',
                'metadata': {},
                'structure': {
                    'headings': [],
                    'tables': [],
                    'images': []
                }
            }
            
            with pdfplumber.open(file_path) as pdf:
                # Extract metadata
                extracted_data['metadata'] = pdf.metadata or {}
                
                for page_num, page in enumerate(pdf.pages):
                    # Extract text
                    page_text = page.extract_text()
                    if page_text:
                        extracted_data['text'] += page_text + "\n\n"
                    
                    # Extract tables with better precision
                    tables = page.extract_tables()
                    for table in tables:
                        extracted_data['structure']['tables'].append({
                            'data': table,
                            'page': page_num + 1
                        })
                        
                        # Add formatted table to text
                        extracted_data['text'] += "\n### Table\n"
                        for row in table:
                            if row:
                                cleaned_row = [str(cell).strip() if cell else "" for cell in row]
                                extracted_data['text'] += " | ".join(cleaned_row) + "\n"
                        extracted_data['text'] += "\n"
            
            return extracted_data
            
        except Exception as e:
            print(f"❌ PDFPlumber extraction failed: {str(e)}")
            return None
    
    def extract_text(self, file_path: str, method: str = "auto") -> Dict[str, Any]:
        """
        Extract text from PDF using specified method
        
        Args:
            file_path: Path to PDF file
            method: "pymupdf", "pdfplumber", or "auto"
            
        Returns:
            Dict containing extracted text and structure information
        """
        print(f"📄 Processing PDF: {os.path.basename(file_path)}")
        
        if method == "auto":
            # Try PyMuPDF first, fallback to pdfplumber
            result = self.extract_with_pymupdf(file_path)
            if result is None:
                result = self.extract_with_pdfplumber(file_path)
        elif method == "pymupdf":
            result = self.extract_with_pymupdf(file_path)
        elif method == "pdfplumber":
            result = self.extract_with_pdfplumber(file_path)
        else:
            raise ValueError(f"Unknown method: {method}")
        
        if result:
            print(f"✅ Extracted {len(result['text'])} characters from PDF")
            print(f"📊 Found {len(result['structure']['tables'])} tables")
            print(f"📋 Found {len(result['structure']['headings'])} headings")
        
        return result


class DOCXProcessor:
    """Extract text from DOCX files while preserving structure"""
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from DOCX file with structure preservation
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dict containing extracted text and structure information
        """
        try:
            print(f"📝 Processing DOCX: {os.path.basename(file_path)}")
            
            doc = Document(file_path)
            extracted_data = {
                'text': '',
                'metadata': {},
                'structure': {
                    'headings': [],
                    'tables': [],
                    'paragraphs': []
                }
            }
            
            # Extract core properties metadata
            if doc.core_properties:
                extracted_data['metadata'] = {
                    'title': doc.core_properties.title,
                    'author': doc.core_properties.author,
                    'subject': doc.core_properties.subject,
                    'created': str(doc.core_properties.created) if doc.core_properties.created else None,
                    'modified': str(doc.core_properties.modified) if doc.core_properties.modified else None,
                }
            
            # Process paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    # Check if paragraph is a heading
                    if para.style.name.startswith('Heading'):
                        level = para.style.name.replace('Heading ', '').strip()
                        extracted_data['structure']['headings'].append({
                            'text': para.text.strip(),
                            'level': level,
                            'style': para.style.name
                        })
                        # Format as markdown heading
                        heading_level = '#' * min(int(level) if level.isdigit() else 1, 6)
                        extracted_data['text'] += f"\n{heading_level} {para.text.strip()}\n\n"
                    else:
                        extracted_data['structure']['paragraphs'].append({
                            'text': para.text.strip(),
                            'style': para.style.name
                        })
                        extracted_data['text'] += para.text.strip() + "\n\n"
            
            # Process tables
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                
                extracted_data['structure']['tables'].append({
                    'data': table_data,
                    'index': table_idx
                })
                
                # Add table to text representation
                extracted_data['text'] += "\n### Table\n"
                for row in table_data:
                    if any(cell.strip() for cell in row):  # Skip empty rows
                        extracted_data['text'] += " | ".join(row) + "\n"
                extracted_data['text'] += "\n"
            
            print(f"✅ Extracted {len(extracted_data['text'])} characters from DOCX")
            print(f"📊 Found {len(extracted_data['structure']['tables'])} tables")
            print(f"📋 Found {len(extracted_data['structure']['headings'])} headings")
            
            return extracted_data
            
        except Exception as e:
            print(f"❌ DOCX extraction failed: {str(e)}")
            return None


class EmailProcessor:
    """Extract text from .eml email files"""
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from EML email file
        
        Args:
            file_path: Path to EML file
            
        Returns:
            Dict containing extracted text and email metadata
        """
        try:
            print(f"📧 Processing EML: {os.path.basename(file_path)}")
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                msg = email.message_from_file(f, policy=policy.default)
            
            extracted_data = {
                'text': '',
                'metadata': {},
                'structure': {
                    'headers': {},
                    'attachments': [],
                    'parts': []
                }
            }
            
            # Extract headers
            headers = {
                'from': msg.get('From', ''),
                'to': msg.get('To', ''),
                'cc': msg.get('Cc', ''),
                'bcc': msg.get('Bcc', ''),
                'subject': msg.get('Subject', ''),
                'date': msg.get('Date', ''),
                'message_id': msg.get('Message-ID', ''),
                'reply_to': msg.get('Reply-To', ''),
            }
            
            extracted_data['metadata'] = headers
            extracted_data['structure']['headers'] = headers
            
            # Format email header for text
            extracted_data['text'] += f"# Email: {headers['subject']}\n\n"
            extracted_data['text'] += f"**From:** {headers['from']}\n"
            extracted_data['text'] += f"**To:** {headers['to']}\n"
            if headers['cc']:
                extracted_data['text'] += f"**CC:** {headers['cc']}\n"
            extracted_data['text'] += f"**Date:** {headers['date']}\n\n"
            
            # Extract body content
            body_text = ""
            
            if msg.is_multipart():
                for part in msg.walk():
                    content_type = part.get_content_type()
                    content_disposition = str(part.get("Content-Disposition", ""))
                    
                    extracted_data['structure']['parts'].append({
                        'content_type': content_type,
                        'disposition': content_disposition
                    })
                    
                    # Skip attachments
                    if "attachment" in content_disposition:
                        filename = part.get_filename()
                        if filename:
                            extracted_data['structure']['attachments'].append(filename)
                        continue
                    
                    # Extract text content
                    if content_type == "text/plain":
                        try:
                            body_text += part.get_content() + "\n\n"
                        except:
                            body_text += str(part.get_payload(decode=True), 'utf-8', errors='ignore') + "\n\n"
                    
                    elif content_type == "text/html":
                        try:
                            html_content = part.get_content()
                            # Convert HTML to plain text
                            soup = BeautifulSoup(html_content, 'html.parser')
                            body_text += soup.get_text() + "\n\n"
                        except:
                            html_content = str(part.get_payload(decode=True), 'utf-8', errors='ignore')
                            soup = BeautifulSoup(html_content, 'html.parser')
                            body_text += soup.get_text() + "\n\n"
            else:
                # Single part message
                content_type = msg.get_content_type()
                if content_type == "text/plain":
                    body_text = msg.get_content()
                elif content_type == "text/html":
                    html_content = msg.get_content()
                    soup = BeautifulSoup(html_content, 'html.parser')
                    body_text = soup.get_text()
            
            extracted_data['text'] += "## Email Body\n\n" + body_text.strip()
            
            # Add attachment information
            if extracted_data['structure']['attachments']:
                extracted_data['text'] += "\n\n## Attachments\n"
                for attachment in extracted_data['structure']['attachments']:
                    extracted_data['text'] += f"- {attachment}\n"
            
            print(f"✅ Extracted {len(extracted_data['text'])} characters from email")
            print(f"📎 Found {len(extracted_data['structure']['attachments'])} attachments")
            
            return extracted_data
            
        except Exception as e:
            print(f"❌ Email extraction failed: {str(e)}")
            return None


class TextChunker:
    """Chunk text into semantically meaningful segments using LangChain"""
    
    def __init__(self, 
                 chunk_size: int = 1000,
                 chunk_overlap: int = 200,
                 separators: Optional[List[str]] = None):
        """
        Initialize text chunker
        
        Args:
            chunk_size: Maximum size of each chunk
            chunk_overlap: Number of characters to overlap between chunks
            separators: Custom separators for splitting (optional)
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Default separators optimized for document structure
        if separators is None:
            separators = [
                "\n\n",      # Paragraph breaks
                "\n",        # Line breaks
                ". ",        # Sentence endings
                "! ",        # Exclamation sentences
                "? ",        # Question sentences
                "; ",        # Semicolon breaks
                ", ",        # Comma breaks
                " ",         # Word breaks
                ""           # Character breaks (last resort)
            ]
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=separators,
            keep_separator=True,
            length_function=len,
        )
    
    def chunk_text(self, text: str, metadata: Dict[str, Any] = None) -> List[Dict[str, Any]]:
        """
        Chunk text into meaningful segments
        
        Args:
            text: Text to chunk
            metadata: Additional metadata to include with chunks
            
        Returns:
            List of chunk dictionaries with text and metadata
        """
        if not text or not text.strip():
            return []
        
        # Split text into chunks
        chunks = self.text_splitter.split_text(text)
        
        # Create chunk objects with metadata
        chunk_objects = []
        for i, chunk_text in enumerate(chunks):
            chunk_obj = {
                'text': chunk_text.strip(),
                'chunk_index': i,
                'chunk_size': len(chunk_text),
                'metadata': metadata or {}
            }
            
            # Add positional metadata
            chunk_obj['metadata'].update({
                'chunk_id': f"chunk_{i}",
                'total_chunks': len(chunks),
                'is_first_chunk': i == 0,
                'is_last_chunk': i == len(chunks) - 1
            })
            
            chunk_objects.append(chunk_obj)
        
        return chunk_objects
    
    def chunk_document(self, document_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Chunk a complete document with structure preservation
        
        Args:
            document_data: Document data from processors
            
        Returns:
            List of chunks with enhanced metadata
        """
        if not document_data or 'text' not in document_data:
            return []
        
        # Prepare base metadata
        base_metadata = document_data.get('metadata', {}).copy()
        base_metadata['has_tables'] = len(document_data.get('structure', {}).get('tables', [])) > 0
        base_metadata['has_headings'] = len(document_data.get('structure', {}).get('headings', [])) > 0
        
        # Chunk the text
        chunks = self.chunk_text(document_data['text'], base_metadata)
        
        # Enhance chunks with structural information
        structure = document_data.get('structure', {})
        
        for chunk in chunks:
            chunk_text = chunk['text'].lower()
            
            # Check if chunk contains tables
            chunk['metadata']['contains_table'] = any(
                'table' in chunk_text or 'data' in chunk_text
                for _ in structure.get('tables', [])
            )
            
            # Check if chunk contains headings
            relevant_headings = []
            for heading in structure.get('headings', []):
                if heading['text'].lower() in chunk_text:
                    relevant_headings.append(heading)
            
            chunk['metadata']['headings'] = relevant_headings
            chunk['metadata']['contains_heading'] = len(relevant_headings) > 0
        
        return chunks
    
    def get_chunk_stats(self, chunks: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Get statistics about the chunks"""
        if not chunks:
            return {}
        
        chunk_sizes = [chunk['chunk_size'] for chunk in chunks]
        
        stats = {
            'total_chunks': len(chunks),
            'total_characters': sum(chunk_sizes),
            'avg_chunk_size': sum(chunk_sizes) / len(chunks),
            'min_chunk_size': min(chunk_sizes),
            'max_chunk_size': max(chunk_sizes),
            'chunks_with_tables': sum(1 for chunk in chunks if chunk['metadata'].get('contains_table', False)),
            'chunks_with_headings': sum(1 for chunk in chunks if chunk['metadata'].get('contains_heading', False))
        }
        
        return stats


class DocumentProcessor:
    """Main document processing pipeline that orchestrates all components"""
    
    def __init__(self, 
                 chunk_size: int = 1000,
                 chunk_overlap: int = 200,
                 temp_dir: Optional[str] = None):
        """
        Initialize the document processor
        
        Args:
            chunk_size: Size of text chunks
            chunk_overlap: Overlap between chunks
            temp_dir: Temporary directory for downloads
        """
        self.downloader = DocumentDownloader()
        self.pdf_processor = PDFProcessor()
        self.docx_processor = DOCXProcessor()
        self.email_processor = EmailProcessor()
        self.text_chunker = TextChunker(chunk_size, chunk_overlap)
        self.temp_dir = temp_dir or tempfile.gettempdir()
        
        # Supported file types
        self.supported_extensions = {
            '.pdf': self.pdf_processor,
            '.docx': self.docx_processor,
            '.doc': self.docx_processor,  # May work with some .doc files
            '.eml': self.email_processor,
        }
    
    def process_file(self, file_path: str, cleanup: bool = True) -> Dict[str, Any]:
        """
        Process a single file through the complete pipeline
        
        Args:
            file_path: Path to file or URL
            cleanup: Whether to delete downloaded files
            
        Returns:
            Dict containing processed document data and chunks
        """
        local_file_path = file_path
        downloaded = False
        
        try:
            # Step 1: Download file if URL
            if self.downloader.is_url(file_path):
                print(f"🌐 Detected URL: {file_path}")
                local_file_path = self.downloader.download_file(file_path)
                downloaded = True
            else:
                print(f"📁 Processing local file: {file_path}")
            
            # Check if file exists
            if not os.path.exists(local_file_path):
                raise FileNotFoundError(f"File not found: {local_file_path}")
            
            # Step 2: Determine file type and processor
            file_extension = Path(local_file_path).suffix.lower()
            
            if file_extension not in self.supported_extensions:
                raise ValueError(f"Unsupported file type: {file_extension}. Supported: {list(self.supported_extensions.keys())}")
            
            processor = self.supported_extensions[file_extension]
            
            # Step 3: Extract text and structure
            print(f"🔍 Extracting content from {file_extension.upper()} file...")
            document_data = processor.extract_text(local_file_path)
            
            if not document_data:
                raise ValueError("Failed to extract content from file")
            
            # Add source information to metadata
            document_data['metadata']['source_file'] = os.path.basename(local_file_path)
            document_data['metadata']['original_path'] = file_path
            document_data['metadata']['file_extension'] = file_extension
            document_data['metadata']['file_size'] = os.path.getsize(local_file_path)
            document_data['metadata']['processed_date'] = str(Path(local_file_path).stat().st_mtime)
            
            # Step 4: Chunk into semantic segments
            print(f"✂️ Chunking document into semantic segments...")
            chunks = self.text_chunker.chunk_document(document_data)
            
            # Step 5: Generate statistics
            chunk_stats = self.text_chunker.get_chunk_stats(chunks)
            
            # Prepare final result
            result = {
                'success': True,
                'source': file_path,
                'local_path': local_file_path,
                'document_data': document_data,
                'chunks': chunks,
                'statistics': {
                    'document': {
                        'total_characters': len(document_data['text']),
                        'file_size_bytes': document_data['metadata']['file_size'],
                        'tables_found': len(document_data['structure'].get('tables', [])),
                        'headings_found': len(document_data['structure'].get('headings', []))
                    },
                    'chunks': chunk_stats
                }
            }
            
            print(f"✅ Processing complete!")
            print(f"📊 Statistics:")
            print(f"   📄 Document: {result['statistics']['document']['total_characters']} chars")
            print(f"   🔤 Chunks: {result['statistics']['chunks']['total_chunks']} pieces")
            print(f"   📋 Tables: {result['statistics']['document']['tables_found']}")
            print(f"   📝 Headings: {result['statistics']['document']['headings_found']}")
            
            return result
            
        except Exception as e:
            print(f"❌ Processing failed: {str(e)}")
            return {
                'success': False,
                'source': file_path,
                'error': str(e),
                'chunks': [],
                'document_data': None
            }
        
        finally:
            # Cleanup downloaded files if requested
            if downloaded and cleanup and os.path.exists(local_file_path):
                try:
                    os.remove(local_file_path)
                    print(f"🗑️ Cleaned up temporary file: {local_file_path}")
                except Exception as e:
                    print(f"⚠️ Could not delete temporary file: {e}")
    
    def process_multiple(self, file_paths: List[str], cleanup: bool = True) -> List[Dict[str, Any]]:
        """
        Process multiple files
        
        Args:
            file_paths: List of file paths or URLs
            cleanup: Whether to delete downloaded files
            
        Returns:
            List of processing results
        """
        results = []
        
        print(f"📂 Processing {len(file_paths)} files...")
        
        for i, file_path in enumerate(file_paths, 1):
            print(f"\n🔄 Processing file {i}/{len(file_paths)}: {file_path}")
            result = self.process_file(file_path, cleanup)
            results.append(result)
        
        # Summary statistics
        successful = sum(1 for r in results if r['success'])
        total_chunks = sum(len(r['chunks']) for r in results if r['success'])
        
        print(f"\n📈 Batch Processing Summary:")
        print(f"   ✅ Successful: {successful}/{len(file_paths)}")
        print(f"   🔤 Total chunks: {total_chunks}")
        
        return results
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats"""
        return list(self.supported_extensions.keys())


# Main execution and initialization
if __name__ == "__main__":
    print("🚀 Document Processor (Ingestion Pipeline)")
    print("=" * 60)
    
    # Initialize the main document processor
    doc_processor = DocumentProcessor(
        chunk_size=1000,
        chunk_overlap=200
    )
    
    # Example usage with a sample file
    result = doc_processor.process_file("https://hackrx.blob.core.windows.net/assets/policy.pdf?sv=2023-01-03&st=2025-07-04T09%3A11%3A24Z&se=2027-07-05T09%3A11%3A00Z&sr=b&sp=r&sig=N4a9OU0w0QXO6AOIBiu4bpl7AXvEZogeT%2FjUHNO7HzQ%3D")
    if result['success']:
        print("📄 Document processed successfully!")
        print(f"Total characters extracted: {len(result['document_data']['text'])}")
        print(f"Total chunks created: {len(result['chunks'])}")
    else:
        print("❌ Document processing failed:", result['error'])